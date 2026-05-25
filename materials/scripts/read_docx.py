from docx import Document
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

filename = sys.argv[1]
d = Document(filename)
for i, para in enumerate(d.paragraphs):
    if para.text.strip():
        print(f'[P{i}] {para.text}')

for ti, table in enumerate(d.tables):
    print(f'\n=== TABLE {ti} ===')
    for row in table.rows:
        print(' | '.join(cell.text for cell in row.cells))
