"""Insert the seven generated figures into research_paper_final.docx.

Each figure is placed just before the heading of the *next* section, with
a centred image and a centred italic caption underneath.
"""
import sys, io, copy
from pathlib import Path
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

MATERIALS = Path(__file__).resolve().parent.parent
FIGURES_DIR = MATERIALS / 'figures'
SRC = str(MATERIALS / 'deliverables' / 'Final_Project_Report' / 'research_paper_final.docx')
DST = SRC

# (image, caption, anchor_heading_text, width_inches)
FIGURES = [
    ('preprocessing_flow.png',
     'Data preprocessing pipeline: from the raw CUAD release to the '
     '4,034-clause training set used by the classifier.',
     '3.2\tSystem Architecture', 6.5),
    ('eda_label_distribution.png',
     'Training-set label distribution across the twelve-category clause '
     'taxonomy. Seven categories are well supported by CUAD; five rely on '
     'the hand-written seed set.',
     '3.2\tSystem Architecture', 6.3),
    ('architecture_diagram.png',
     'Two-service system architecture: a Next.js web application on Vercel '
     'handles UI, authentication and persistence; a FastAPI service on '
     'Hugging Face Spaces handles all NLP inference.',
     '3.3\tMethodology / Algorithm / Model Description', 5.0),
    ('model_architecture.png',
     'Six-stage NLP pipeline executed by the ML service for every document. '
     'Stages 3–6 (classify, NER, rewrite, risk-flag) run once per clause.',
     '3.4\tTools, Technologies and Frameworks Used', 5.5),
    ('model_selection.png',
     'Candidate models considered for clause classification and intent '
     'rewriting, with the rationale for the chosen baseline and the planned '
     'transformer upgrade.',
     '3.4\tTools, Technologies and Frameworks Used', 6.4),
    ('metrics_formulae.png',
     'Evaluation metrics used in this work: precision, recall, F₁, '
     'accuracy, macro- and weighted-average F₁, and the TF–IDF feature '
     'weighting used by the classifier.',
     'CHAPTER 4\tEXPERIMENTATION AND RESULTS', 6.3),
    ('results_chart.png',
     'Per-class held-out performance of the fine-tuned Legal-BERT classifier '
     '(freak3123/legal-bert-cuad-v1). Overall accuracy is 0.912, '
     'weighted-average F₁ is 0.915, and macro-average F₁ is 0.902 across all '
     'twelve clause categories.',
     '4.4.2\tEnd-to-End Pipeline Output on a Sample Contract', 6.5),
    ('comparison_baseline_bert.png',
     'Headline metric improvement from baseline to fine-tuned: TF-IDF + '
     'Logistic Regression macro F₁ of 0.61 against Legal-BERT macro F₁ of '
     '0.902 on the same held-out 419-clause test set.',
     '4.4.2\tEnd-to-End Pipeline Output on a Sample Contract', 5.5),
    ('confusion_matrix.png',
     'Confusion matrix for the fine-tuned Legal-BERT classifier on the held-out '
     '419-clause test set. Off-diagonal mass is concentrated on OTHER ↔ '
     'INTELLECTUAL_PROPERTY and WARRANTY boundary cases.',
     '4.4.3\tAutomated Test Suite', 6.0),
    ('flan_t5_rouge_curves.png',
     'Validation ROUGE for the fine-tuned FLAN-T5-small clause simplifier '
     '(freak3123/flan-t5-simplify-v1) across ten training epochs. Final '
     'scores: ROUGE-1 = 0.394, ROUGE-2 = 0.231, ROUGE-L = 0.352.',
     '4.4.3\tAutomated Test Suite', 6.0),
    ('legal_bert_training_curves.png',
     'Legal-BERT fine-tuning trajectory across the four training epochs on a '
     'Kaggle T4 GPU. Training and validation losses fall monotonically while '
     'validation macro F1 climbs from 0.512 at epoch 1 to a best of 0.885 at '
     'epoch 3; the trainer saves the epoch-3 checkpoint as the best model.',
     '4.4.3\tAutomated Test Suite', 6.0),
    ('rewrite_examples.png',
     'Side-by-side legal-to-plain-English rewrites produced by the fine-tuned '
     'FLAN-T5-small model on three representative clause types (automatic '
     'renewal, user-content licence, data sharing). All outputs are verbatim '
     'model generations, with no manual editing.',
     '4.4.3\tAutomated Test Suite', 6.5),
]


def find_paragraph(doc, text):
    """Find the first paragraph whose text equals the given string."""
    # Tabs in heading text are stored as \t in python-docx output.
    for p in doc.paragraphs:
        if p.text == text:
            return p
    # fallback: case- and whitespace-insensitive match
    norm = lambda s: ''.join(s.split()).lower()
    target = norm(text)
    for p in doc.paragraphs:
        if norm(p.text) == target:
            return p
    return None


def insert_paragraph_before(target_p, doc):
    """Insert a new empty paragraph immediately before target_p; return it."""
    new_p = copy.deepcopy(target_p._p)
    # strip all children so we start clean
    for child in list(new_p):
        new_p.remove(child)
    # clear pPr (paragraph properties) so we don't inherit Heading style
    target_p._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, target_p._parent)


def main():
    doc = Document(SRC)

    figure_no = 0
    for img, caption, anchor, width in FIGURES:
        figure_no += 1
        target = find_paragraph(doc, anchor)
        if target is None:
            print(f'!! anchor not found for {img}: {anchor!r}')
            continue

        # 1. picture paragraph
        p_img = insert_paragraph_before(target, doc)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(FIGURES_DIR / img), width=Inches(width))

        # 2. caption paragraph
        p_cap = insert_paragraph_before(target, doc)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_cap.style = doc.styles['Caption']
        except KeyError:
            pass
        r = p_cap.add_run(f'Figure {figure_no}: {caption}')
        r.italic = True
        r.font.size = Pt(10)

        print(f'inserted Figure {figure_no} ({img}) before {anchor!r}')

    doc.save(DST)
    print(f'wrote {DST}')


if __name__ == '__main__':
    main()
