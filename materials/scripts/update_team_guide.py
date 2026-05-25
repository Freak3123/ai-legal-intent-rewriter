# -*- coding: utf-8 -*-
"""Bring project_team_guide.docx up to date with Phase 3 reality.

The original guide was written at end-of-Phase-2 (TF-IDF baseline, template
rewriter, no OCR, regex-only NER). Phase 3 is now complete and the guide's
"Incomplete features", "Known limitations" and Q&A sections need updating so
the senior-lecturer review doesn't read stale fact claims.

Run from anywhere:
    python materials/scripts/update_team_guide.py
"""
import io
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
PATH = REPO_ROOT / "project_team_guide.docx"

REPLACEMENTS: list[tuple[str, str]] = [
    # Section 6 framing
    (
        "This section is an honest status check, drawn from docs/HANDOVER.md and the 'STUB' / 'Phase' comments in the code. The project is at the end of 'Phase 2': the pipeline works end to end, but several parts are deliberately simple placeholders waiting for 'Phase 3'.",
        "This section is an honest status check. The project is at the end of Phase 3: every pipeline stage now runs against either a real fine-tuned model or a robust rule layer, both transformer checkpoints are pushed to the Hugging Face Hub, and the system is ready for Phase 4 deployment.",
    ),
    # "Incomplete features" entries — these are NOT incomplete anymore
    (
        "Clause rewriting is only a template. rewrite.py does not yet use a real model — it picks a fixed opening sentence per category and does simple word swaps. The planned FLAN-T5-small model is not trained or wired in.",
        "Clause rewriting now uses the fine-tuned FLAN-T5-small model freak3123/flan-t5-simplify-v1 (validation ROUGE-L = 0.352). The template-based simplifier is kept as an automatic fallback that activates if the Hub model is unreachable.",
    ),
    (
        "The classifier is a baseline, not a neural model. The planned Legal-BERT model has not been fine-tuned yet.",
        "The classifier is a fine-tuned Legal-BERT model (freak3123/legal-bert-cuad-v1, macro F1 = 0.902 on the held-out test set). The TF-IDF + Logistic Regression baseline (macro F1 0.61) is retained as a runtime fallback for cold-start scenarios.",
    ),
    (
        "Scanned PDFs are not supported. ingest.py's PDF function raises 'NotImplementedError'. Server-side PDF parsing (PyMuPDF) and OCR (Tesseract) are planned. The UI already detects a scanned PDF and shows an 'OCR not yet wired' message.",
        "Scanned PDFs are supported end to end. ingest.py uses PyMuPDF for the text layer and falls back to Tesseract OCR when the layer is empty. The web app's upload page detects scanned PDFs client-side and POSTs them to a new /v1/analyze/pdf endpoint.",
    ),
    (
        "Entity recognition is regex-only. ner.py uses deliberately strict regular expressions, so it favours precision and will miss some valid entities. The planned upgrade is a spaCy statistical model.",
        "Entity recognition uses a two-layer hybrid: restrictive regex for all six entity types (precision layer) plus spaCy en_core_web_sm for DATE / AMOUNT / PARTY recall. The regex layer wins on overlap, so spaCy adds recall without sacrificing precision.",
    ),
    # Known limitations
    (
        "Five clause categories are weak. PAYMENT, CONFIDENTIALITY, DISPUTE_RESOLUTION, DEFINITIONS and RENEWAL have only 6 training examples each (from the seed CSV), because the CUAD dataset does not cover them. Their cross-validation scores are very low. The model still predicts them at runtime, but accuracy on them is unreliable. Fix: add more examples to data/clauses_seed.csv and re-train.",
        "Two clause categories — WARRANTY (F1 0.735) and DISPUTE_RESOLUTION (F1 0.750) — show high recall but lower precision, meaning the model occasionally over-predicts them. The seed CSV was expanded to 36–39 examples each for the five previously sparse categories (PAYMENT, CONFIDENTIALITY, DISPUTE_RESOLUTION, DEFINITIONS, RENEWAL), and all five now score F1 0.75–1.00. Further seed-set expansion for WARRANTY and DISPUTE_RESOLUTION is the simplest remaining mitigation.",
    ),
    (
        "Not deployed yet. The apps run locally. Deployment to Vercel and Hugging Face Spaces is planned ('Phase 4').",
        "Deployment to Vercel (web) and Hugging Face Spaces (ML) is the final remaining step. The deployment runbook is in docs/deploy-guide.md and the Dockerfile bakes in the spaCy model so the Space's cold start is cached.",
    ),
    # Section 7 Q2 — TF-IDF baseline question
    (
        "Q2. Why is the classifier only TF-IDF + Logistic Regression?",
        "Q2. Why fine-tune Legal-BERT rather than building a classical baseline?",
    ),
    (
        "It is a deliberate, honest baseline. It is fast, tiny (2 MB), runs on any CPU, and gives a measured number for the planned Legal-BERT model to beat. On a stratified cross-validation it reaches 0.89 overall accuracy, with F1-scores of 0.84–0.98 on the seven well-supported categories. Starting with a strong, well-understood baseline before a heavyweight model is good machine-learning practice.",
        "Both approaches were built. A TF-IDF + Logistic Regression baseline was trained first to give a measurable target (macro F1 = 0.61 on the held-out test set). It is kept in the codebase as a runtime fallback. The production classifier is then a fine-tuned nlpaueb/legal-bert-base-uncased model (4 epochs, T4 GPU, class-balanced cross-entropy), which reaches macro F1 = 0.902 on the same test set — a 48% relative improvement over the baseline. Building both lets us quantify the lift the transformer actually delivers, which is good machine-learning practice.",
    ),
    # Section 7 Q3 — macro F1 question
    (
        "Q3. Your macro F1 is only 0.56 — is the model bad?",
        "Q3. How did you reach macro F1 = 0.902 across all twelve categories?",
    ),
    (
        "No — it is a data problem, not a model problem. Seven categories have hundreds of CUAD examples and score 0.81–0.98. Five categories have only six examples each, so cross-validation cannot learn them and their scores collapse to near zero, which drags the macro average down. The fix is more training data, and it is already the top priority in the future-work plan.",
        "Two changes together. First, the seed CSV was expanded from 72 to 227 hand-curated examples, adding 30–33 examples each for the five categories CUAD does not cover (PAYMENT, CONFIDENTIALITY, DISPUTE_RESOLUTION, DEFINITIONS, RENEWAL). Second, the classifier was upgraded from TF-IDF + LogReg to a fine-tuned Legal-BERT with class-balanced cross-entropy loss. Either change alone would have helped; together they raise macro F1 from 0.61 (baseline + expanded seed) to 0.902 (transformer + expanded seed), with every one of the 12 categories scoring F1 ≥ 0.735 on the held-out test set.",
    ),
    # Section 7 Q5 — test suite
    (
        "There is an automated test suite of 34 pytest tests covering segmentation, classification, entity recognition, the health endpoint and the full analyze endpoint — all passing. The classifier is evaluated with stratified cross-validation. The whole pipeline has been run end to end on sample contracts and produces correct, structured output.",
        "An automated pytest suite of 34 tests covers segmentation, classification, entity recognition, the health endpoint and the full analyze endpoint — all passing. The classifier is evaluated on a held-out 419-clause test set (10% stratified split, seed 42 — the same split the Kaggle training notebook used), reproducible via apps/ml/scripts/eval_classifier.py. Per-class precision/recall/F1 and a confusion matrix are in docs/evaluation.md.",
    ),
    # Section 7 Q7 — what would you do with more time
    (
        "In priority order: expand the training data for the five weak categories; fine-tune Legal-BERT for classification and FLAN-T5-small for rewriting; add server-side PDF parsing and OCR; upgrade entity recognition to a spaCy statistical model; deploy both services; and run a user study with real participants.",
        "In priority order: deploy both services to Vercel + Hugging Face Spaces; build a manually annotated test set for the segmentation, NER and risk-flagging modules and report their entity-level F1 / precision / recall; expand the seed CSV for the two lower-precision categories (WARRANTY, DISPUTE_RESOLUTION); run a user study with five to ten non-expert participants; and add a continuous-improvement loop that captures in-product corrections and feeds them back into periodic re-training.",
    ),
    # Section 7.4 "Working well"
    (
        "The end-to-end pipeline runs reliably and is covered by 34 passing tests.",
        "The end-to-end pipeline runs reliably with fine-tuned Legal-BERT and FLAN-T5-small in production, covered by 34 passing tests.",
    ),
    (
        "Clause classification is strong on every well-supported category (F1 0.84–0.98).",
        "Clause classification is strong across all twelve categories on the held-out test set (macro F1 = 0.902; every category ≥ 0.735).",
    ),
    # Section 7.4 "Could improve" — items that are now done
    (
        "Train the real neural models (Legal-BERT, FLAN-T5-small) to replace the baseline classifier and the template rewriter.",
        "Build a manually annotated test set to report quantitative segmentation, NER and risk-flagging metrics (the modules are presently validated functionally by the test suite, qualitatively by worked examples).",
    ),
    (
        "Add training examples for the five weak categories.",
        "Add 20–30 more seed examples each for WARRANTY and DISPUTE_RESOLUTION to lift their precision.",
    ),
    (
        "Add scanned-PDF (OCR) support.",
        "Run a small user study (5–10 non-expert participants) on real contracts and report readability and trust ratings.",
    ),
    (
        "Move long analyses to a background job and deploy the system publicly.",
        "Deploy publicly (Vercel + Hugging Face Spaces). The deployment runbook is in docs/deploy-guide.md.",
    ),
    # Final-note honesty paragraph
    (
        "Final note for reviewers: the project is honest about its own status. It is a working, tested, well-architected system at the end of its second development phase, with a clear and realistic plan for the remaining work.",
        "Final note for reviewers: the project is honest about its own status. It is a working, tested, well-architected system at the end of its third development phase — every model is fine-tuned and on the Hub, every pipeline module is real (no stubs), the test suite is green, and the only remaining work is public deployment plus quantitative evaluation of the rule-based modules.",
    ),
]


def replace_in_paragraph(p: Paragraph, old: str, new: str) -> bool:
    text = "".join(r.text for r in p.runs)
    if old not in text:
        return False
    new_text = text.replace(old, new)
    if not p.runs:
        return False
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text
    return True


def main() -> int:
    doc = Document(str(PATH))
    paragraphs = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    changes = 0
    misses: list[str] = []
    for old, new in REPLACEMENTS:
        hit = False
        for p in paragraphs:
            if replace_in_paragraph(p, old, new):
                changes += 1
                hit = True
                break
        if not hit:
            misses.append(old[:80] + "...")

    doc.save(str(PATH))
    print(f"applied {changes}/{len(REPLACEMENTS)} replacements; saved {PATH.name}")
    if misses:
        print("\nMISSED replacements (text not found verbatim):")
        for m in misses:
            print(f"  - {m}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
