# -*- coding: utf-8 -*-
"""Update project_team_guide.docx now that Phase 4 (deployment) is complete.

Both services are live:
  * Web:  https://ai-legal-intent-rewriter-web.vercel.app
  * ML:   https://freak3123-legal-rewriter-ml.hf.space
"""
import io
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
PATH = REPO_ROOT / "project_team_guide.docx"

REPLACEMENTS: list[tuple[str, str]] = [
    # Phase framing — bump from Phase 3 to Phase 4 complete
    (
        "This section is an honest status check. The project is at the end of Phase 3: every pipeline stage now runs against either a real fine-tuned model or a robust rule layer, both transformer checkpoints are pushed to the Hugging Face Hub, and the system is ready for Phase 4 deployment.",
        "This section is an honest status check. The project is at the end of Phase 4: every pipeline stage runs against either a real fine-tuned model or a robust rule layer, both transformer checkpoints are pushed to the Hugging Face Hub, and both services are deployed publicly — the Next.js web app on Vercel at https://ai-legal-intent-rewriter-web.vercel.app and the FastAPI ML service on Hugging Face Spaces at https://freak3123-legal-rewriter-ml.hf.space. The only remaining work is quantitative evaluation of the rule-based modules and the planned user study.",
    ),
    # Deployment item — flip from "remaining step" to "done"
    (
        "Deployment to Vercel (web) and Hugging Face Spaces (ML) is the final remaining step. The deployment runbook is in docs/deploy-guide.md and the Dockerfile bakes in the spaCy model so the Space's cold start is cached.",
        "Both services are deployed and publicly reachable. The Next.js web app is live at https://ai-legal-intent-rewriter-web.vercel.app (Vercel Hobby tier, MongoDB Atlas M0 free tier for data). The FastAPI ML service is live at https://freak3123-legal-rewriter-ml.hf.space (Hugging Face Spaces, Docker SDK, free CPU basic). The Space pulls both fine-tuned models from the Hub on first cold start (~30 seconds) and then serves sub-second responses thereafter. The deployment runbook is in docs/deploy-guide.md.",
    ),
    # Cold-start mitigation — was "concern", now "known + documented"
    (
        "Analyses run synchronously. The /api/analyses route waits for the ML service inside the web request. A very long contract could hit Vercel's 10-second timeout once deployed. Fix: a background job queue, or call the ML service directly from the browser.",
        "Analyses run synchronously: the /api/analyses route waits for the ML service inside the web request. The Vercel route's maxDuration is set to 60 seconds, which covers the cold-start case (Space spins up Legal-BERT + FLAN-T5 in ~30 seconds) plus a 10-page contract on the free CPU tier. Very long contracts could still time out — the long-term fix is a background job queue or direct browser-to-Space calls with a short-lived JWT.",
    ),
    # Q7 future-work — drop the "deploy" item since it's done, push other items up
    (
        "In priority order: deploy both services to Vercel + Hugging Face Spaces; build a manually annotated test set for the segmentation, NER and risk-flagging modules and report their entity-level F1 / precision / recall; expand the seed CSV for the two lower-precision categories (WARRANTY, DISPUTE_RESOLUTION); run a user study with five to ten non-expert participants; and add a continuous-improvement loop that captures in-product corrections and feeds them back into periodic re-training.",
        "In priority order: build a manually annotated test set for the segmentation, NER and risk-flagging modules and report their entity-level F1 / precision / recall; expand the seed CSV for the two lower-precision categories (WARRANTY, DISPUTE_RESOLUTION); run a user study with five to ten non-expert participants; pre-warm the Space with a scheduled cron ping so the cold-start cost disappears for live demos; and add a continuous-improvement loop that captures in-product corrections and feeds them back into periodic re-training.",
    ),
    # §7.4 Could-improve — drop the deploy item, replace with something else useful
    (
        "Deploy publicly (Vercel + Hugging Face Spaces). The deployment runbook is in docs/deploy-guide.md.",
        "Eliminate the ~30-second cold-start on the Space by adding a scheduled keep-alive ping or upgrading to the always-on HF Space tier ($9/month).",
    ),
    # §7.4 Working well — add a deployment line
    (
        "The architecture is clean, well-documented and genuinely deployable.",
        "The architecture is clean, well-documented, and now genuinely deployed — public URLs on Vercel and Hugging Face Spaces serve the full pipeline end to end.",
    ),
    # Final note — bump from phase 3 to phase 4
    (
        "Final note for reviewers: the project is honest about its own status. It is a working, tested, well-architected system at the end of its third development phase — every model is fine-tuned and on the Hub, every pipeline module is real (no stubs), the test suite is green, and the only remaining work is public deployment plus quantitative evaluation of the rule-based modules.",
        "Final note for reviewers: the project is honest about its own status. It is a working, tested, well-architected system at the end of its fourth development phase. Every model is fine-tuned and on the Hugging Face Hub; every pipeline module is real (no stubs); the 34-test pytest suite is green; the web app is live on Vercel and the ML service is live on Hugging Face Spaces; and the remaining work is quantitative evaluation of the rule-based modules and a small user study.",
    ),
]


def replace_in_paragraph(p: Paragraph, old: str, new: str) -> bool:
    text = "".join(r.text for r in p.runs)
    if old not in text:
        return False
    new_text = text.replace(old, new)
    if not p.runs:
        return False
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text
    return True


def main() -> int:
    doc = Document(str(PATH))
    paragraphs = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    changes = 0
    misses: list[str] = []
    for old, new in REPLACEMENTS:
        hit = False
        for p in paragraphs:
            if replace_in_paragraph(p, old, new):
                changes += 1
                hit = True
                break
        if not hit:
            misses.append(old[:80] + "...")

    doc.save(str(PATH))
    print(f"applied {changes}/{len(REPLACEMENTS)} Phase-4 replacements; saved {PATH.name}")
    if misses:
        print("\nMISSED replacements:")
        for m in misses:
            print(f"  - {m}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
